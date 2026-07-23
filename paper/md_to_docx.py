"""
md_to_docx.py - Convert paper.md to JUTIF template-compatible DOCX

Usage: .venv/bin/python paper/md_to_docx.py paper/paper.md paper/paper.docx
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def parse_md(md_path: Path):
    """Parse markdown into structured sections."""
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    elements = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # YAML frontmatter / title
        if line.startswith('# '):
            elements.append(('title', line[2:].strip()))
            i += 1
        elif line.startswith('## '):
            elements.append(('h1', line[3:].strip()))
            i += 1
        elif line.startswith('### '):
            elements.append(('h2', line[4:].strip()))
            i += 1
        elif line.startswith('#### '):
            elements.append(('h3', line[5:].strip()))
            i += 1
        elif line.startswith('|') and '|' in line[1:]:
            # Markdown table
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                # Skip separator row (|---|---|)
                if not all(set(c) <= set('-:') for c in row):
                    table_rows.append(row)
                i += 1
            if table_rows:
                elements.append(('table', table_rows))
        elif line.startswith('---'):
            elements.append(('hr', ''))
            i += 1
        elif line.startswith('```'):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing
            elements.append(('code', '\n'.join(code_lines)))
        elif line.startswith('1. ') or re.match(r'^\d+\.\s', line):
            # Numbered list
            while i < len(lines) and (re.match(r'^\d+\.\s', lines[i]) or (lines[i].startswith('   ') and lines[i].strip())):
                elements.append(('ol', lines[i].strip()))
                i += 1
        elif line.strip().startswith('- '):
            # Bullet list
            while i < len(lines) and (lines[i].strip().startswith('- ') or (lines[i].startswith('  ') and lines[i].strip())):
                elements.append(('ul', lines[i].strip()))
                i += 1
        elif line.strip().startswith('**') and '**' in line.strip()[2:]:
            # Bold paragraph
            elements.append(('p_bold', line.strip()))
            i += 1
        elif line.strip() == '':
            elements.append(('blank', ''))
            i += 1
        else:
            # Regular paragraph
            elements.append(('p', line.strip()))
            i += 1

    return elements


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown paper to JUTIF template DOCX."""
    doc = Document()
    elements = parse_md(md_path)

    # Set base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # Set margins per JUTIF: 2.5cm left/top, 2cm right/bottom
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)

    for kind, content in elements:
        if kind == 'title':
            # Title: centered, bold, 16pt
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.size = Pt(16)
            run.bold = True
            run.font.name = 'Times New Roman'

        elif kind == 'h1':
            # Section header: bold, 10pt, centered or left
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(content.upper())
            run.font.size = Pt(10)
            run.bold = True
            run.font.name = 'Times New Roman'

        elif kind == 'h2':
            # Subsection: bold italic, 10pt
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(10)
            run.bold = True
            run.italic = True

        elif kind == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(10)
            run.bold = True

        elif kind == 'p_bold':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Strip surrounding **
            text = content.strip('*').strip()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)

        elif kind == 'p':
            if not content:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

        elif kind == 'blank':
            doc.add_paragraph()

        elif kind == 'table':
            rows = content
            if not rows:
                continue
            n_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Times New Roman'
                            if r_idx == 0:
                                run.bold = True

        elif kind == 'code':
            # Code block — monospace, smaller
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        elif kind == 'ul':
            # Bullet list item
            p = doc.add_paragraph(style='List Bullet')
            text = content.lstrip('- ').strip()
            run = p.add_run(text)
            run.font.size = Pt(10)

        elif kind == 'ol':
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', content)
            run = p.add_run(text)
            run.font.size = Pt(10)

        elif kind == 'hr':
            doc.add_paragraph('—' * 20)

    doc.save(str(docx_path))
    print(f"Saved {docx_path}")
    print(f"  Elements: {len(elements)}")
    print(f"  Tables: {sum(1 for e in elements if e[0] == 'table')}")
    print(f"  Headers: {sum(1 for e in elements if e[0] in ('h1', 'h2', 'h3'))}")


if __name__ == '__main__':
    md = Path(sys.argv[1])
    docx = Path(sys.argv[2])
    md_to_docx(md, docx)
