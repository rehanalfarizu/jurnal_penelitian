#!/usr/bin/env python3
"""Build the UAS research proposal as a fully formatted Word document."""

from __future__ import annotations

import json
import os
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", "/tmp/matplotlib-uasp")
os.environ.setdefault("XDG_CACHE_HOME", "/tmp/uasp-cache")
os.environ.setdefault("MPLBACKEND", "Agg")
Path("/tmp/matplotlib-uasp").mkdir(parents=True, exist_ok=True)
Path("/tmp/uasp-cache").mkdir(parents=True, exist_ok=True)

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
FIGURES = PAPER / "figures"
OUTPUT = PAPER / "Proposal_UAS_Metodologi_Penelitian_Rehan_Alfarizu.docx"
LOGO = FIGURES / "logo_amikom_template.png"
FINAL_RESULTS = ROOT / "results" / "final"

TITLE = (
    "EVALUASI KINERJA DIGITAL TWIN EDGE–CLOUD MULTISKALA "
    "UNTUK MONITORING ENERGI DAN OKUPANSI"
)
AUTHOR = "REHAN ALFARIZU"
NIM = "[NIM BELUM DITEMUKAN]"
PROGRAM = "S1 INFORMATIKA"

PURPLE = "#5B2C83"
BLUE = "#2B6CB0"
CYAN = "#2C7A7B"
GREEN = "#2F855A"
ORANGE = "#C05621"
GRAY = "#4A5568"
LIGHT = "#F7FAFC"


def box(ax, xy, width, height, text, color, fontsize=9, linewidth=1.6):
    patch = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.025,rounding_size=0.07",
        facecolor="white",
        edgecolor=color,
        linewidth=linewidth,
    )
    ax.add_patch(patch)
    ax.text(
        xy[0] + width / 2,
        xy[1] + height / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        color="#1A202C",
        wrap=True,
    )
    return patch


def arrow(ax, start, end, color=GRAY, rad=0.0, style="-|>"):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle=style,
            mutation_scale=13,
            linewidth=1.5,
            color=color,
            connectionstyle=f"arc3,rad={rad}",
        )
    )


def save_architecture_figure(path: Path):
    fig, ax = plt.subplots(figsize=(14, 7.6), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title(
        "Arsitektur Digital Twin Edge–Cloud dan Visualisasi Multiskala",
        fontsize=16,
        weight="bold",
        pad=14,
        color="#1A202C",
    )

    layers = [
        (0.25, "LAPISAN FISIK", PURPLE),
        (3.6, "EDGE", BLUE),
        (7.0, "CLOUD TEREMULASI", ORANGE),
        (10.4, "DIGITAL TWIN MULTISKALA", GREEN),
    ]
    for x, label, color in layers:
        ax.add_patch(
            FancyBboxPatch(
                (x, 0.45),
                3.05,
                6.75,
                boxstyle="round,pad=0.03,rounding_size=0.09",
                facecolor=LIGHT,
                edgecolor=color,
                linewidth=2.0,
            )
        )
        ax.text(
            x + 1.525,
            6.83,
            label,
            ha="center",
            va="center",
            fontsize=10,
            weight="bold",
            color=color,
        )

    box(ax, (0.62, 5.25), 2.3, 0.85, "Sensor suhu & kelembapan\nDHT11", PURPLE)
    box(ax, (0.62, 3.85), 2.3, 0.85, "Sensor tegangan ZMPT101B\n& arus SCT013", PURPLE)
    box(ax, (0.62, 2.45), 2.3, 0.85, "Okupansi / jumlah orang\n(alur terpisah)", PURPLE)
    box(ax, (0.62, 1.05), 2.3, 0.85, "Trace historis satu gateway\n92.160 baris", PURPLE)

    box(ax, (3.98, 5.25), 2.3, 0.85, "Akuisisi & validasi\ntelemetry", BLUE)
    box(ax, (3.98, 3.85), 2.3, 0.85, "Validasi skema, timestamp\n& status sensor", BLUE)
    box(ax, (3.98, 2.45), 2.3, 0.85, "Energi legacy Wh,\nokupansi & provenance", BLUE)
    box(ax, (3.98, 1.05), 2.3, 0.85, "Routing lokal\nnormal → edge", BLUE)

    box(ax, (7.38, 5.25), 2.3, 0.85, "Baseline cloud-only\nterkonfigurasi", ORANGE)
    box(ax, (7.38, 3.85), 2.3, 0.85, "Jalur anomali / invalid\n→ cloud", ORANGE)
    box(ax, (7.38, 2.45), 2.3, 0.85, "Profil jaringan terkonfigurasi\nlatensi, jitter, drop", ORANGE)
    box(ax, (7.38, 1.05), 2.3, 0.85, "Audit, agregasi &\nprovenance", ORANGE)

    box(ax, (10.78, 5.25), 2.3, 0.85, "Kontrak telemetry JSON\n& replay API", GREEN)
    box(ax, (10.78, 3.85), 2.3, 0.85, "LoD-A · Tapak\nEPSG:4326", GREEN)
    box(ax, (10.78, 2.45), 2.3, 0.85, "LoD-B · Bangunan\nenergi + okupansi", GREEN)
    box(ax, (10.78, 1.05), 2.3, 0.85, "LoD-C · Indoor 3D\nVue + Babylon.js", GREEN)

    arrow(ax, (2.92, 4.27), (3.98, 5.62), PURPLE, rad=-0.15)
    arrow(ax, (6.28, 2.87), (7.38, 4.27), ORANGE, rad=-0.15)
    arrow(ax, (6.28, 1.47), (10.78, 5.65), BLUE, rad=-0.15)
    arrow(ax, (9.68, 4.27), (10.78, 5.65), ORANGE, rad=0.15)
    arrow(ax, (11.93, 5.25), (11.93, 4.70), GREEN)
    arrow(ax, (11.93, 3.85), (11.93, 3.30), GREEN)
    arrow(ax, (11.93, 2.45), (11.93, 1.90), GREEN)

    ax.text(
        7,
        0.1,
        "Catatan: latensi jaringan/cloud dievaluasi melalui emulasi terkontrol; "
        "bukan pengukuran Azure produksi.",
        ha="center",
        fontsize=9,
        color=GRAY,
        style="italic",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_provenance_figure(path: Path):
    fig, ax = plt.subplots(figsize=(14, 7.3), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    ax.set_title(
        "Provenance dan Pemanfaatan Replay Data Historis",
        fontsize=16,
        weight="bold",
        pad=14,
        color="#1A202C",
    )
    ax.text(0.45, 6.65, "SUMBER DATA HISTORIS", weight="bold", color=BLUE, fontsize=11)
    ax.text(0.45, 3.25, "WORKLOAD REPLAY DAN EVALUASI ARSITEKTUR", weight="bold", color=ORANGE, fontsize=11)

    y1, y2 = 5.35, 1.95
    items_top = [
        (0.45, "Trace asli\n92.160 baris"),
        (3.0, "Audit skema, timestamp,\nnilai nol & missing"),
        (5.55, "Verifikasi daya legacy\nV×I dan satuan"),
        (8.1, "Pemetaan source_row_id\n& checksum"),
        (10.65, "Baseline karakteristik\ntrace historis"),
    ]
    items_bottom = [
        (0.45, "Workload replay legacy\n2.027.520 baris"),
        (3.0, "Rekonstruksi 22 blok\nhistorical replay"),
        (5.55, "Replay terurut & sampel\nmerata 5.000 pesan"),
        (8.1, "Validasi, serialisasi,\nrouting & jaringan"),
        (10.65, "P50/P95/P99,\nthroughput & deadline"),
    ]
    for x, text in items_top:
        box(ax, (x, y1), 2.05, 0.9, text, BLUE, 9)
    for x, text in items_bottom:
        box(ax, (x, y2), 2.05, 0.9, text, ORANGE, 9)
    for items, y, color in [(items_top, y1, BLUE), (items_bottom, y2, ORANGE)]:
        for idx in range(len(items) - 1):
            arrow(ax, (items[idx][0] + 2.05, y + 0.45), (items[idx + 1][0], y + 0.45), color)

    box(ax, (11.15, 3.63), 2.45, 0.95, "Telemetry berprovenance\n→ API → tiga skala", GREEN, 9.5, 2.0)
    arrow(ax, (11.68, 5.35), (12.0, 4.58), GREEN)
    arrow(ax, (11.68, 2.85), (12.0, 3.63), GREEN)
    ax.text(
        7,
        0.62,
        "Jumlah 2.027.520 menyatakan volume workload replay, bukan jumlah observasi lapangan independen.",
        ha="center",
        fontsize=10,
        weight="bold",
        color="#9C4221",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_research_flow_figure(path: Path):
    fig, ax = plt.subplots(figsize=(10.2, 14.7), dpi=200)
    ax.set_xlim(0, 10.4)
    ax.set_ylim(0, 15)
    ax.axis("off")

    border = "#27364A"
    arrow_color = "#34465C"
    text_color = "#172033"
    phase_fill = ("#FFFFFF", "#F7F9FC", "#FFFFFF")
    phase_bounds = ((8.95, 14.30), (4.45, 8.95), (0.20, 4.45))
    phase_labels = (
        "PERSIAPAN DAN AUDIT DATA",
        "EVALUASI ARSITEKTUR EDGE–CLOUD",
        "INTEGRASI DIGITAL TWIN DAN PELAPORAN",
    )

    def rect(x, y, width, height, facecolor="white", linewidth=1.15, zorder=1):
        patch = FancyBboxPatch(
            (x, y),
            width,
            height,
            boxstyle="square,pad=0",
            facecolor=facecolor,
            edgecolor=border,
            linewidth=linewidth,
            zorder=zorder,
        )
        ax.add_patch(patch)
        return patch

    def text_box(x, y, width, height, label, fontsize=8.6, linewidth=1.1):
        rect(x, y, width, height, "white", linewidth, 3)
        ax.text(
            x + width / 2,
            y + height / 2,
            label,
            ha="center",
            va="center",
            multialignment="center",
            fontsize=fontsize,
            linespacing=1.16,
            color=text_color,
            zorder=4,
        )

    def group_box(y, title, labels):
        group_x, group_w, group_h = 1.35, 8.55, 1.34
        rect(group_x, y, group_w, group_h, "#FBFCFE", 1.15, 2)
        ax.text(
            group_x + group_w / 2,
            y + 1.10,
            title,
            ha="center",
            va="center",
            fontsize=8.6,
            weight="bold",
            color=text_color,
            zorder=4,
        )
        sub_width = 2.35
        for x, label in zip((1.65, 4.45, 7.25), labels):
            text_box(x, y + 0.18, sub_width, 0.70, label, 7.7, 1.0)

    def down_arrow(upper_bottom, lower_top):
        ax.add_patch(
            FancyArrowPatch(
                (5.625, upper_bottom),
                (5.625, lower_top),
                arrowstyle="-|>",
                mutation_scale=11,
                linewidth=1.15,
                color=arrow_color,
                zorder=5,
            )
        )

    # Kerangka dan judul mengikuti struktur diagram alur penelitian formal.
    rect(0.10, 14.30, 10.15, 0.55, "#E8EDF4", 1.4)
    ax.text(
        5.175,
        14.575,
        "ALUR PENELITIAN",
        ha="center",
        va="center",
        fontsize=12.2,
        weight="bold",
        color=text_color,
    )
    for index, ((bottom, top), label) in enumerate(zip(phase_bounds, phase_labels)):
        rect(0.10, bottom, 10.15, top - bottom, phase_fill[index], 1.2)
        rect(0.10, bottom, 0.78, top - bottom, "#EEF2F7", 1.0, 2)
        ax.text(
            0.49,
            (bottom + top) / 2,
            label,
            ha="center",
            va="center",
            rotation=90,
            fontsize=7.8,
            weight="bold",
            color=text_color,
            zorder=4,
        )

    # Fase 1 — persiapan, audit, dan pembentukan workload replay.
    text_box(2.55, 13.52, 6.15, 0.55, "Penetapan ruang lingkup dan batas klaim")
    text_box(2.55, 12.73, 6.15, 0.55, "Akuisisi data historis")
    group_box(
        11.08,
        "Audit dan Persiapan Data",
        (
            "Trace asli\n92.160 baris",
            "Workload replay\n2.027.520 baris",
            "Audit skema,\ntimestamp, dan lineage",
        ),
    )
    text_box(
        2.55,
        10.15,
        6.15,
        0.56,
        "Rekonstruksi 22 blok replay dan pemetaan provenance",
        8.4,
    )
    text_box(
        2.55,
        9.23,
        6.15,
        0.56,
        "Integrasi energi legacy dan status okupansi",
        8.4,
    )
    down_arrow(13.52, 13.28)
    down_arrow(12.73, 12.42)
    down_arrow(11.08, 10.71)
    down_arrow(10.15, 9.79)

    # Fase 2 — pengukuran komparatif edge–cloud.
    text_box(
        2.55,
        8.15,
        6.15,
        0.55,
        "Pemilihan 5.000 sampel merata lintas blok",
        8.5,
    )
    group_box(
        6.55,
        "Pemrosesan Pesan dan Pembanding",
        (
            "Validasi lokal\ndan serialisasi",
            "Routing selektif\nedge–cloud",
            "Baseline\ncloud-only",
        ),
    )
    text_box(
        2.55,
        5.65,
        6.15,
        0.55,
        "Emulasi jaringan: latensi, jitter, dan packet drop",
        8.4,
    )
    text_box(
        2.55,
        4.72,
        6.15,
        0.55,
        "Pengukuran P50/P95/P99, throughput, payload, dan deadline miss",
        8.15,
    )
    down_arrow(8.95, 8.70)
    down_arrow(8.15, 7.89)
    down_arrow(6.55, 6.20)
    down_arrow(5.65, 5.27)

    # Fase 3 — kontrak data, visualisasi multiskala, pengujian, dan pelaporan.
    text_box(
        2.55,
        3.72,
        6.15,
        0.55,
        "Validasi kontrak telemetry dan replay API",
        8.5,
    )
    group_box(
        2.05,
        "Visualisasi Digital Twin Multiskala",
        (
            "LoD-A\nTapak",
            "LoD-B\nBangunan",
            "LoD-C\nIndoor 3D",
        ),
    )
    text_box(
        2.55,
        1.20,
        6.15,
        0.55,
        "Pengujian schema, API, frontend, dan notebook",
        8.4,
    )
    text_box(
        2.55,
        0.37,
        6.15,
        0.55,
        "Analisis hasil, ancaman validitas, dan penyusunan laporan",
        8.3,
    )
    down_arrow(4.45, 4.27)
    down_arrow(3.72, 3.39)
    down_arrow(2.05, 1.75)
    down_arrow(1.20, 0.92)

    fig.subplots_adjust(left=0.025, right=0.975, top=0.985, bottom=0.015)
    fig.savefig(path, bbox_inches="tight", pad_inches=0.05, facecolor="white")
    plt.close(fig)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, text, end):
        run._r.append(child)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, line=1.5):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Cm(1.27) if indent else Cm(0)


def add_body(doc, text, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold_lead=None):
    p = doc.add_paragraph()
    format_paragraph(p, align=align, indent=indent)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_list(doc, items, ordered=True, level=0):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.left_indent = Cm(0.64 + level * 0.64)
        fmt.first_line_indent = Cm(-0.64)
        fmt.line_spacing = 1.5
        fmt.space_after = Pt(4)
        marker = f"{idx}. " if ordered else "• "
        r1 = p.add_run(marker)
        set_run_font(r1)
        r2 = p.add_run(item)
        set_run_font(r2)


def chapter(doc, number, title):
    if len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"BAB {number}\n{title.upper()}")
    set_run_font(r, bold=True)


def subheading(doc, number, title, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"{number} {title}")
    set_run_font(r, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text))
    return p


def add_figure(doc, path, caption, width_cm=14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_paragraph = add_caption(doc, caption)
    caption_paragraph.paragraph_format.keep_with_next = False


def add_equation(doc, expression, number):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [12.2, 1.8]
    for column, width in zip(table.columns, widths):
        column.width = Cm(width)
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    middle = table.cell(0, 0).paragraphs[0]
    middle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    middle.paragraph_format.first_line_indent = Cm(0)
    middle.paragraph_format.space_after = Pt(0)
    equation_size = 11 if len(expression) > 50 else 12
    set_run_font(
        middle.add_run(expression),
        size=equation_size,
        italic=True,
    )
    right_cell = table.cell(0, 1)
    no_wrap = OxmlElement("w:noWrap")
    right_cell._tc.get_or_add_tcPr().append(no_wrap)
    right = right_cell.paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.first_line_indent = Cm(0)
    right.paragraph_format.space_after = Pt(0)
    set_run_font(right.add_run(f"({number})"))


def add_simple_table(doc, headers, rows, widths, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, (text, width) in enumerate(zip(headers, widths)):
        cell = header.cells[i]
        set_cell_width(cell, width)
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=font_size, bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, (text, width) in enumerate(zip(row_data, widths)):
            cell = row.cells[i]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 4 else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(str(text)), size=font_size)
    doc.add_paragraph()
    return table


def build_document():
    FIGURES.mkdir(parents=True, exist_ok=True)
    architecture = FIGURES / "arsitektur_edge_cloud_digital_twin.png"
    provenance = FIGURES / "pemisahan_peran_data.png"
    research_flow = FIGURES / "alur_penelitian_final.png"
    save_architecture_figure(architecture)
    save_provenance_figure(provenance)
    save_research_flow_figure(research_flow)
    benchmark_results = json.loads(
        (FINAL_RESULTS / "benchmark_metrics.json").read_text(encoding="utf-8")
    )

    doc = Document()
    props = doc.core_properties
    props.title = TITLE.title()
    props.author = AUTHOR.title()
    props.subject = "Proposal UAS Metodologi Penelitian"
    props.keywords = (
        "digital twin, edge-cloud, energi, okupansi, geospasial, indoor, "
        "multiskala, near real-time, replay data historis"
    )

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.space_after = Pt(6)

    if "Caption UAS" not in [s.name for s in doc.styles]:
        caption_style = doc.styles.add_style("Caption UAS", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style.font.size = Pt(12)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(TITLE), size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("PROPOSAL PENELITIAN"), size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("untuk memenuhi tugas Ujian Akhir Semester\nMata Kuliah Metodologi Penelitian"), size=12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(5))
        p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("disusun oleh"), size=12)
    p.paragraph_format.space_after = Pt(6)

    for value in (AUTHOR, NIM):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(value), size=12, bold=True)
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.line_spacing = 1.5
    set_run_font(
        p.add_run(
            f"PROGRAM STUDI {PROGRAM}\n"
            "FAKULTAS ILMU KOMPUTER\n"
            "UNIVERSITAS AMIKOM YOGYAKARTA\n"
            "YOGYAKARTA\n"
            "2026"
        ),
        size=12,
        bold=True,
    )

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Cm(21)
    body_section.page_height = Cm(29.7)
    body_section.left_margin = Cm(4)
    body_section.right_margin = Cm(3)
    body_section.top_margin = Cm(3)
    body_section.bottom_margin = Cm(3)
    body_section.footer.is_linked_to_previous = False
    add_page_number(body_section.footer.paragraphs[0])
    set_page_number_start(body_section, 1)

    # BAB I
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("BAB I\nPENDAHULUAN"), bold=True)

    subheading(doc, "1.1", "Latar Belakang")
    background = [
        (
            "Pengelolaan energi bangunan memerlukan telemetri yang cepat dan dapat ditelusuri agar "
            "perubahan beban dan kondisi ruang dapat dipantau. Internet of Things (IoT) mendukung "
            "pengumpulan data kelistrikan, lingkungan, dan okupansi, tetapi pemrosesan yang seluruhnya "
            "bergantung pada cloud dapat menambah latensi dan lalu lintas komunikasi. Arsitektur "
            "edge–cloud membagi validasi dan routing berlatensi rendah di edge serta penyimpanan dan "
            "pemrosesan lebih lanjut di cloud [1]–[4]. Evaluasinya perlu memisahkan latensi komputasi "
            "yang diukur dari latensi jaringan yang dikonfigurasi atau diemulasi."
        ),
        (
            "Digital Twin menghubungkan telemetri dengan representasi virtual bangunan. Integrasi "
            "BIM–GIS dan Web-3D memungkinkan kondisi energi serta okupansi ditampilkan pada skala "
            "tapak, bangunan, dan indoor [5]–[7], [14]–[18]. Namun, data historis penelitian hanya "
            "berasal dari satu gateway dan firmware lama menghitung daya sebagai tegangan dikalikan "
            "arus tanpa faktor daya. Oleh karena itu, provenance dan batas klaim diperlukan agar "
            "nilai tersebut tidak disalahartikan sebagai pengukuran daya aktif terkalibrasi."
        ),
        (
            "Penelitian ini mengevaluasi arsitektur pemantauan menggunakan replay data historis. "
            "Workbook 92.160 baris menjadi trace rujukan, sedangkan CSV 2.027.520 baris diposisikan "
            "sebagai 22 pengulangan deterministik dari blok turunan yang diaudit, bukan observasi "
            "lapangan independen atau salinan mentah workbook. Workload tersebut digunakan untuk "
            "menilai latensi, throughput, routing, deadline near real-time, energi legacy, okupansi, "
            "pembanding cloud-only, dan konsistensi visualisasi Digital Twin multiskala."
        ),
    ]
    for para in background:
        add_body(doc, para)

    subheading(doc, "1.2", "Rumusan Masalah")
    add_body(
        doc,
        "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah sebagai berikut.",
    )
    add_list(
        doc,
        [
            "Bagaimana data historis asli dan workload replay dipersiapkan dengan provenance yang dapat ditelusuri untuk monitoring energi dan okupansi?",
            "Bagaimana karakteristik latensi, throughput, routing, dan pemenuhan batas near real-time pada arsitektur edge–cloud dibandingkan baseline cloud-only terkonfigurasi?",
            "Bagaimana energi turunan dari daya legacy dan status okupansi disajikan tanpa disalahartikan sebagai pengukuran meter atau validasi sensor lapangan?",
            "Bagaimana satu kontrak telemetry divisualisasikan secara konsisten pada skala tapak geospasial, bangunan, dan indoor?",
        ],
    )

    subheading(doc, "1.3", "Batasan Masalah")
    add_body(doc, "Batasan masalah yang diterapkan agar penelitian tetap terukur adalah sebagai berikut.")
    add_list(
        doc,
        [
            "Objek fisik dibatasi pada prototipe bangunan cerdas yang direpresentasikan oleh satu gateway historis RASPBERRY_PI_GATEWAY_001.",
            "Trace asli berupa workbook 92.160 baris dengan rentang aktual sekitar 19–23 Mei 2026 dan digunakan untuk audit, karakterisasi, serta rujukan provenance.",
            "Daya pada firmware lama merupakan hasil tegangan × arus tanpa faktor daya; karena itu dilaporkan sebagai daya legacy atau proksi daya semu, bukan daya aktif independen.",
            "Energi dihitung sebagai integral trapesium proksi daya legacy terhadap timestamp sumber dengan gap maksimum 10 detik; hasilnya bukan energi aktif terkalibrasi.",
            "CSV workload legacy 2.027.520 baris merupakan 22 replay dari trace asli dan dilaporkan sebagai volume workload, bukan 2 juta observasi lapangan.",
            "Penelitian tidak melatih atau menilai model prediksi daya dan tidak melaporkan metrik akurasi seperti MAE, RMSE, atau R².",
            "Pemrosesan edge dibatasi pada validasi skema, verifikasi nilai, penambahan provenance, serialisasi, dan routing pesan.",
            "Komputasi lokal diukur pada mesin eksperimen, sedangkan latensi jaringan/cloud merupakan emulasi terkonfigurasi dan bukan pengukuran Azure produksi.",
            "Near real-time didefinisikan sebagai penyelesaian jalur pemrosesan sebelum deadline operasional terkonfigurasi 3,5 detik—pembulatan median interval trace 3,5251918 detik, bukan interval publish firmware—disertai pelaporan P50, P95, P99, dan deadline miss.",
            "Digital Twin dibatasi pada fungsi monitoring dan sinkronisasi satu arah; penelitian tidak mengklaim kontrol otonom atau closed-loop.",
            "Visualisasi menerapkan LoD aplikatif proyek: LoD-A tapak, LoD-B bangunan, dan LoD-C indoor 3D. Koordinat belum diverifikasi survei dan kepatuhan LoD geometrik CityGML, IndoorGML, IFC, atau 3D Tiles belum dievaluasi.",
            "Latensi render peramban belum dimasukkan ke latency end-to-end dan dinyatakan sebagai keterbatasan.",
        ],
    )

    subheading(doc, "1.4", "Tujuan Penelitian")
    add_body(doc, "Tujuan yang hendak dicapai dalam penelitian ini adalah sebagai berikut.")
    add_list(
        doc,
        [
            "Membangun pipeline replay data historis yang mempertahankan keterurutan, identitas baris sumber, blok replay, dan provenance telemetry.",
            "Mengukur karakteristik edge–cloud dan membandingkannya dengan baseline cloud-only menggunakan workload serta profil jaringan terkonfigurasi yang sama.",
            "Mengintegrasikan energi legacy per siklus, status okupansi, rute, freshness, dan provenance ke kontrak telemetry serta replay API.",
            "Menyajikan payload yang sama pada visual tapak geospasial, bangunan, dan indoor secara konsisten.",
        ],
    )

    subheading(doc, "1.5", "Manfaat Penelitian")
    add_body(
        doc,
        "Secara teoritis, penelitian ini memberikan rancangan evaluasi pemantauan yang memisahkan karakteristik sumber data, volume workload replay, latensi komputasi terukur, dan latensi jaringan teremulasi. Rancangan tersebut membantu mencegah pengulangan data historis disalahartikan sebagai observasi independen.",
    )
    add_body(
        doc,
        "Secara praktis, penelitian ini menghasilkan pipeline yang dapat direplikasi untuk mengaudit trace, merekonstruksi blok replay, menghitung energi-proksi, memantau okupansi, membandingkan edge–cloud dengan cloud-only, serta menyajikan telemetry berprovenance pada visual tapak–bangunan–indoor. Hasilnya dapat membantu pengembang menilai kesiapan arsitektur monitoring tanpa memerlukan sensor aktif baru.",
    )

    subheading(doc, "1.6", "Sistematika Penulisan")
    add_body(doc, "Sistematika penulisan penelitian direncanakan sebagai berikut.")
    add_list(
        doc,
        [
            "BAB I PENDAHULUAN memuat latar belakang, rumusan masalah, batasan masalah, tujuan, manfaat, dan sistematika penulisan.",
            "BAB II TINJAUAN PUSTAKA membahas bangunan cerdas, energi–okupansi, edge–cloud, near real-time, replay, provenance, Digital Twin, dan integrasi geospasial–indoor.",
            "BAB III METODE PENELITIAN menjelaskan objek, sumber dan posisi data, alur penelitian, variabel, rancangan evaluasi, alat, serta bahan.",
            "BAB IV HASIL DAN PEMBAHASAN menyajikan audit data, hasil benchmark, pembanding cloud-only, energi–okupansi, dan konsistensi telemetry pada tiga skala.",
            "BAB V PENUTUP memuat kesimpulan yang menjawab rumusan masalah serta saran pengembangan dan validasi lapangan berikutnya.",
        ],
    )

    # BAB II
    chapter(doc, "II", "Tinjauan Pustaka")
    subheading(doc, "2.1", "Studi Literatur")
    add_body(
        doc,
        "Studi mengenai manajemen energi bangunan menunjukkan pergeseran dari monitoring terpusat menuju sistem IoT yang dipadukan dengan analitik dan komputasi terdistribusi. Shahid, Shahid, dan Irfan [1] mengidentifikasi IoT, simulasi, serta integrasi sumber energi sebagai unsur penting sistem manajemen energi, tetapi juga menekankan masalah kualitas data, interoperabilitas, dan belum seragamnya evaluasi.",
    )
    add_body(
        doc,
        "Pada tingkat arsitektur, Himeur dkk. [2] membandingkan edge-only, cloud-only, dan hybrid edge–cloud untuk analisis energi bangunan. Arsitektur hibrida memberi kompromi antara kecepatan, biaya cloud, dan latensi komunikasi. Kajian Trigka dan Dritsas [3] serta Rojek dkk. [4] juga menempatkan edge sebagai pemroses lokal yang responsif dan cloud sebagai penyedia kapasitas serta penyimpanan yang skalabel. Namun, generalisasi kinerja tetap bergantung pada perangkat, jaringan, workload, dan cara pengukuran.",
    )
    add_body(
        doc,
        "Digital Twin pada bangunan digunakan untuk monitoring waktu nyata, pemodelan prediktif, integrasi energi, dan pemeliharaan [5]. Sinthamrongruk, Dahal, dan Harnpornchai [6] menunjukkan bahwa Web3D berbasis peramban dapat mengintegrasikan sensor IoT dan visualisasi tiga dimensi dengan kebutuhan deployment yang ringan. Wang dkk. [7] menggunakan MQTT, WebSocket, dan Three.js untuk visualisasi model bangunan secara waktu nyata. Kedua penelitian tersebut menguatkan peran Web-3D, tetapi belum berfokus pada provenance dan keterlacakan workload replay historis.",
    )
    add_body(
        doc,
        "Chen, Chen, dan Huang [14] menunjukkan visualisasi BIM dalam lingkungan indoor–outdoor multiskala, sedangkan Herle dkk. [15] membahas interoperabilitas model geospasial dan BIM. Clausen dkk. [16] serta Smirnov dan Re Cecconi [17] menghubungkan Digital Twin dengan energi, kenyamanan, dan okupansi. Walczyk dan Ożadowicz [18] memetakan integrasi BIM, Digital Twin, otomasi bangunan, serta IoT terdistribusi. Literatur ini menjadi dasar tiga skala visual dan hubungan energi–okupansi, tanpa berarti prototipe telah memenuhi standar BIM–GIS formal.",
    )
    add_body(
        doc,
        "Penelitian prediksi energi menggunakan Transformer [8], LSTM–GNN [9], serta informasi okupansi multimodal [10]. Pendekatan tersebut memerlukan target dan variasi data yang memadai. Penelitian ini tidak mengadopsi pemodelan prediktif karena trace hanya berasal dari satu gateway dan kolom daya merupakan hasil V×I firmware. Fokus dialihkan pada integritas telemetry, performa pemrosesan replay, dan konsistensi pembaruan Digital Twin.",
    )
    add_body(
        doc,
        "Berdasarkan studi tersebut, celah penelitian terletak pada evaluasi terpadu yang mempertahankan provenance sumber, membandingkan edge–cloud dan cloud-only menggunakan replay workload, menurunkan indikator energi secara transparan, serta memeriksa konsistensi telemetry energi–okupansi pada visual tapak–bangunan–indoor. Penelitian ini membedakan jumlah pesan replay dari observasi lapangan dan membatasi klaim geospasial maupun metrologi.",
    )

    add_caption(doc, "Tabel 2.1 Keaslian Penelitian")
    lit_rows = [
        ["1", "Hybrid edge–cloud untuk efisiensi energi bangunan", "Himeur dkk.", "2022", "Arsitektur hibrida memberi kompromi pemrosesan, biaya cloud, dan latensi.", "Penelitian ini membandingkan edge–cloud dengan baseline cloud-only pada replay berprovenance."],
        ["2", "Edge and Cloud Computing in Smart Cities", "Trigka dan Dritsas", "2025", "Merangkum arsitektur, alokasi sumber daya, real-time analytics, dan tantangan keamanan.", "Penelitian ini melakukan benchmark terukur dan emulasi jaringan dengan batas klaim eksplisit."],
        ["3", "Digital Twin untuk efisiensi energi bangunan", "Sghiri dkk.", "2025", "Digital Twin mendukung monitoring, prediksi, integrasi energi, dan pemeliharaan.", "Penelitian ini mengimplementasikan kontrak energi–okupansi satu arah pada studi satu gateway."],
        ["4", "Web3D Digital Twin untuk monitoring ESG", "Sinthamrongruk dkk.", "2026", "Web3D browser-native mengintegrasikan sensor IoT dan visualisasi waktu nyata.", "Penelitian ini menambahkan replay historis, routing edge–cloud, dan provenance setiap pesan."],
        ["5", "Manajemen model bangunan terintegrasi MQTT", "Wang dkk.", "2025", "MQTT–WebSocket dan Three.js mendukung visualisasi sensor pada model bangunan.", "Penelitian ini menambahkan replay API serta provenance trace sumber dan blok replay."],
        ["6", "Transformer untuk prediksi konsumsi energi multibangunan", "Moveh dkk.", "2025", "Transformer memodelkan korelasi temporal dan lintas bangunan.", "Penelitian ini tidak melakukan prediksi; fokusnya pemantauan satu gateway dan auditabilitas replay."],
        ["7", "Forecasting energi berbasis okupansi LSTM–GNN", "Suharto dkk.", "2026", "Okupansi dan dependensi spasial meningkatkan peramalan pada dataset publik.", "Penelitian ini tidak melakukan forecasting; fokusnya pemantauan nilai historis dan kinerja arsitektur."],
        ["8", "Fusi multimodal untuk deteksi okupansi", "Sun", "2024", "Transformer memadukan citra dan audio untuk meningkatkan deteksi okupansi.", "Penelitian tidak menambah fusi multimodal; nilai okupansi hanya ditampilkan sesuai provenance historis."],
        ["9", "Visualisasi BIM indoor–outdoor multiskala", "Chen dkk.", "2021", "Metode berbasis semantik mengatur visualisasi BIM dalam virtual globe dan ruang indoor.", "Penelitian ini menerapkan tiga skala aplikasi tanpa mengklaim konversi BIM–GIS formal."],
        ["10", "Occupancy-aware Digital Twin", "Smirnov dan Re Cecconi", "2026", "Digital Twin menghubungkan energi, okupansi, IEQ, dan visualisasi pada 26 ruang.", "Penelitian ini memakai jumlah orang legacy tanpa meminjam akurasi inferensi artikel."],
    ]
    add_simple_table(
        doc,
        ["No", "Judul penelitian", "Nama penulis", "Tahun", "Hasil penelitian", "Perbandingan penelitian"],
        lit_rows,
        [0.7, 2.7, 2.0, 1.0, 3.8, 3.8],
        font_size=8,
    )

    subheading(doc, "2.2", "Dasar Teori")
    subheading(doc, "2.2.1", "Bangunan Cerdas dan Telemetry", level=2)
    add_body(
        doc,
        "Bangunan cerdas menggunakan sensor, konektivitas, komputasi, dan aplikasi untuk mengamati serta mendukung pengelolaan kondisi bangunan. Telemetry adalah data yang dikirim perangkat secara periodik, misalnya suhu, kelembapan, tegangan, arus, daya, status sensor, timestamp, dan identitas perangkat. Kualitas telemetry dipengaruhi oleh interval sampling, kuantisasi, dropout, packet loss, sinkronisasi waktu, dan aturan firmware [1], [4].",
    )

    subheading(doc, "2.2.2", "Daya, Energi Legacy, dan Okupansi", level=2)
    add_body(
        doc,
        "Daya aktif pada sistem arus bolak-balik satu fase secara umum dipengaruhi tegangan RMS, arus RMS, dan faktor daya. Firmware historis tidak mengukur faktor daya sehingga hanya menghitung hasil V×I dan menyimpannya dengan satuan watt. Nilai tersebut diperlakukan sebagai daya legacy atau proksi daya semu. Sistem memantau nilai itu sesuai sumbernya tanpa mengubahnya menjadi klaim daya aktif independen.",
    )
    add_equation(doc, "S = V × I", "2-1")
    add_body(doc, "dengan S adalah daya semu, V adalah tegangan RMS, dan I adalah arus RMS.", indent=False)
    add_body(
        doc,
        "Energi legacy dihitung dengan integrasi trapesium terhadap timestamp sumber. Interval hanya dipakai bila positif dan tidak lebih dari 10 detik; nilai kumulatif di-reset pada awal setiap siklus replay. Karena masukan tetap proksi V×I, hasil Wh bukan pembacaan meter energi aktif. Okupansi dilaporkan dari jumlah orang legacy sebagai status terisi atau kosong dan tidak diperlakukan sebagai ground truth baru.",
    )
    add_equation(doc, "Eᵢ = ((Pᵢ₋₁ + Pᵢ) / 2) × Δt / 3600", "2-2")

    subheading(doc, "2.2.3", "Arsitektur Edge–Cloud", level=2)
    add_body(
        doc,
        "Edge computing menempatkan validasi, praproses, dan routing dekat sumber data untuk mengurangi perjalanan data dan latensi. Cloud menyediakan penyimpanan, agregasi, serta sumber daya yang lebih besar. Arsitektur hibrida mengatur jalur lokal untuk pesan normal dan jalur cloud untuk data invalid atau anomali [2]–[4]. Pemilihan jalur perlu disertai aturan routing serta provenance agar hasil dapat diaudit.",
    )
    add_figure(doc, architecture, "Gambar 2.1 Arsitektur Digital Twin edge–cloud dan visualisasi multiskala", 14.0)

    subheading(doc, "2.2.4", "Near Real-Time", level=2)
    add_body(
        doc,
        "Near real-time berarti hasil tersedia cukup cepat terhadap kebutuhan aplikasi, tetapi tidak menjanjikan batas deterministik seperti hard real-time. Penelitian ini menggunakan deadline operasional terkonfigurasi 3,5 detik, yaitu pembulatan median interval trace asli 3,5251918 detik dan bukan interval publish nominal firmware. Karakteristik distribusi dilaporkan melalui persentil P50, P95, dan P99, sedangkan deadline miss menunjukkan proporsi pesan yang melebihi batas tersebut.",
    )

    subheading(doc, "2.2.5", "Digital Twin Geospasial–Indoor Multiskala", level=2)
    add_body(
        doc,
        "Digital Twin menghubungkan representasi digital dengan data aset fisik untuk monitoring, analisis, dan interaksi. Integrasi BIM–GIS menjembatani konteks tapak dan detail bangunan, sedangkan Web-3D menyajikan ruang indoor melalui peramban [5]–[7], [14], [15]. Antarmuka penelitian memakai satu kontrak telemetry untuk tiga tingkat detail aplikatif: LoD-A tapak geospasial EPSG:4326, LoD-B ringkasan bangunan, dan LoD-C indoor Babylon. Detail semantik dan visual meningkat antartingkat, dengan perpindahan melalui pilihan tampilan manual. Kontrak memuat energi legacy, okupansi, provenance, rute, timestamp, dan latensi. Klasifikasi LoD ini didefinisikan proyek dan belum menyatakan kepatuhan LoD geometrik CityGML, IndoorGML, IFC, atau 3D Tiles; ruang lingkupnya tetap monitoring satu arah.",
    )

    subheading(doc, "2.2.6", "Replay Data Historis dan Provenance", level=2)
    add_body(
        doc,
        "Replay data historis adalah pemutaran kembali pesan yang telah direkam dengan urutan dan interval yang dipertahankan atau dikendalikan. Setiap pesan harus membawa identitas sumber, indeks baris, blok replay, timestamp sumber, timestamp replay, dan source_type. Pengulangan meningkatkan volume workload untuk pengujian sistem, tetapi tidak menambah jumlah observasi lapangan independen.",
    )
    add_figure(doc, provenance, "Gambar 2.2 Provenance dan pemanfaatan replay data historis", 14.0)

    subheading(doc, "2.2.7", "Validasi, Routing, dan Freshness Telemetry", level=2)
    add_body(
        doc,
        "Validasi pesan memeriksa skema, timestamp, tipe data, nilai wajib, dan status kelistrikan. Routing mempertahankan pesan normal pada edge dan mengirim pesan invalid atau melampaui ambang daya legacy ke jalur cloud terkonfigurasi. Freshness benchmark adalah proksi durasi pemrosesan sampai payload tersedia; umur kalender data historis dilaporkan terpisah melalui timestamp sumber.",
    )

    subheading(doc, "2.2.8", "Metrik Evaluasi", level=2)
    add_body(
        doc,
        "Evaluasi arsitektur menggunakan latensi pemrosesan dan serialisasi, latensi end-to-end, throughput, ukuran payload, rasio routing, drop, deadline miss, dan konsistensi pesan. Edge–cloud dibandingkan dengan baseline cloud-only menggunakan pesan, payload, seed, dan draw jaringan yang sama. Persentil P50 menunjukkan median, sedangkan P95 dan P99 menggambarkan ekor distribusi.",
    )
    add_equation(doc, "Throughput = jumlah pesan selesai / waktu pemrosesan", "2-3")
    add_equation(doc, "Deadline miss rate = jumlah pesan terlambat / jumlah pesan terkirim", "2-4")
    add_equation(doc, "Freshness proxy = latensi komputasi + serialisasi + jaringan sesuai rute", "2-5")
    add_body(
        doc,
        "Latensi komputasi lokal dibedakan dari tambahan latensi jaringan. Nilai jaringan selalu diberi label configured/emulated. Konsistensi telemetry diperiksa melalui kelengkapan field, urutan source_row_id, kesesuaian nilai yang dikirim dan ditampilkan, serta jumlah pesan diterima atau gagal.",
    )

    # BAB III
    chapter(doc, "III", "Metode Penelitian")
    subheading(doc, "3.1", "Objek Penelitian")
    add_body(
        doc,
        "Objek penelitian adalah prototipe Digital Twin edge–cloud untuk monitoring energi dan okupansi near real-time pada bangunan cerdas dengan visual tapak–bangunan–indoor. Sistem merepresentasikan satu gateway historis yang menerima telemetry suhu, kelembapan, tegangan, arus, daya legacy, dan jumlah orang. Komponen fisik historis meliputi ESP32, DHT11, ZMPT101B, dan SCT013, sedangkan komponen virtual menggunakan layanan replay, kontrak JSON, Vue, SVG geospasial, serta Babylon.js.",
    )
    add_body(
        doc,
        "Penelitian berbentuk evaluasi sistem menggunakan replay data historis. Penelitian bukan eksperimen operasional pada banyak bangunan, bukan pengukuran Raspberry Pi baru, dan bukan validasi public cloud. Unit evaluasi adalah pesan telemetry yang dapat ditelusuri kembali ke source_row_id dan replay_block_id pada workload historis.",
    )
    add_caption(doc, "Tabel 3.1 Komponen objek penelitian")
    object_rows = [
        ["Trace rujukan", "Workbook sensor asli", "92.160 baris; satu gateway; sekitar empat hari"],
        ["Workload replay", "CSV turunan legacy", "2.027.520 baris atau 22 replay; bukan data independen"],
        ["Provenance", "Identitas sumber dan replay", "source_row_id, replay_block_id, timestamp, source_type"],
        ["Edge–cloud", "Validasi, serialisasi, routing, jaringan", "Pemrosesan lokal terukur; jaringan/cloud teremulasi"],
        ["Digital Twin", "Tapak, bangunan, indoor dan telemetry", "EPSG:4326, Vue, Babylon.js, JSON schema, replay API"],
    ]
    add_simple_table(doc, ["Komponen", "Bentuk", "Ruang lingkup"], object_rows, [3.3, 4.0, 6.7], font_size=9)

    subheading(doc, "3.2", "Alur Penelitian")
    add_body(
        doc,
        "Penelitian dimulai dengan menetapkan ruang lingkup dan batas klaim. Workbook asli diaudit untuk memahami skema, definisi daya, cadence, nilai nol, dan karakteristik sumber. Workload turunan legacy direkonstruksi sebagai 22 blok replay dengan provenance eksplisit. Energi-proksi diintegrasikan pada timestamp sumber, okupansi dipetakan, pesan diproses melalui edge–cloud dan baseline cloud-only, lalu payload yang sama ditampilkan pada skala tapak, bangunan, dan indoor.",
    )
    add_figure(doc, research_flow, "Gambar 3.1 Alur penelitian", 13.4)

    stages = [
        ("Penetapan ruang lingkup", "menentukan pertanyaan penelitian, definisi operasional, metrik, dan batas klaim."),
        ("Audit trace asli", "memeriksa jumlah baris, rentang timestamp, interval sampling, perangkat, nilai nol, nilai hilang, kuantil, serta checksum."),
        ("Verifikasi daya dan energi legacy", "memastikan hubungan V×I, mengintegrasikan Wh per siklus, dan mendokumentasikan ketiadaan faktor daya serta meter independen."),
        ("Karakterisasi trace", "menghitung cadence, kuantil, proporsi nol, nilai hilang, autokorelasi, dan distribusi status telemetry."),
        ("Audit workload replay", "membuktikan hubungan 22 replay, memeriksa keterurutan timestamp, dan membedakan volume workload dari observasi independen."),
        ("Rekonstruksi provenance", "menambahkan source_type, replay_block_id, source_row_id, timestamp sumber, dan timestamp replay."),
        ("Persiapan replay", "mempertahankan urutan pesan serta mengambil 5.000 posisi merata lintas seluruh blok untuk benchmark per-pesan."),
        ("Pemrosesan edge", "memvalidasi skema dan nilai, membawa energi–okupansi, menambahkan metadata, serta melakukan serialisasi JSON."),
        ("Routing edge–cloud", "mengirim pesan normal melalui edge dan pesan invalid atau anomali melalui profil cloud terkonfigurasi."),
        ("Benchmark arsitektur", "mengukur latency, throughput, payload, routing, drop, deadline miss, freshness, serta pembanding cloud-only."),
        ("Integrasi Digital Twin", "mengirim energi, okupansi, rute, timestamp, dan provenance melalui replay API ke visual tapak–bangunan–indoor."),
        ("Analisis validitas", "membandingkan hasil dengan tujuan serta menyatakan keterbatasan data, perangkat, jaringan, dan render peramban."),
    ]
    for index, (name, explanation) in enumerate(stages, start=1):
        add_body(doc, f"{index}. {name}: {explanation}", indent=False, bold_lead=f"{index}. {name}:")

    subheading(doc, "3.3", "Alat dan Bahan")
    doc.add_page_break()
    subheading(doc, "3.3.1", "Data Penelitian", level=2)
    data_rows = [
        ["Trace historis asli", "sensor_data_export_2026-05-17_to_2026-05-23.xlsx", "92.160", "Audit, karakterisasi, dan rujukan provenance"],
        ["Workload replay historis", "sensor_data.csv; 22 replay trace turunan", "2.027.520", "Benchmark edge–cloud/cloud-only dan visual multiskala"],
    ]
    add_caption(doc, "Tabel 3.2 Posisi data penelitian")
    add_simple_table(doc, ["Jenis", "Sumber/provenance", "Jumlah baris", "Penggunaan"], data_rows, [3.0, 5.0, 2.0, 4.0], font_size=9)
    add_body(
        doc,
        "Data asli diperoleh dari ekspor telemetry gateway historis. Workload turunan legacy tidak diperlakukan sebagai tambahan sampel statistik karena payload mengulang pola trace sumber. Setiap baris workload dipetakan kembali ke blok dan indeks sumber agar jumlah replay, urutan, dan provenance dapat diaudit.",
    )

    subheading(doc, "3.3.2", "Alat dan Instrumen", level=2)
    tool_rows = [
        ["Komputer eksperimen", "macOS 15.7, arsitektur x86_64", "Menjalankan audit, benchmark, replay, dan dashboard"],
        ["Python", "3.11.9", "Audit, rekonstruksi provenance, benchmark, dan pelaporan"],
        ["pandas / NumPy", "Sesuai environment eksperimen", "Pengolahan tabel, deret waktu, dan perhitungan numerik"],
        ["JSON dan high-resolution clock", "Schema validation dan monotonic timer", "Validasi pesan, serialisasi, dan pengukuran latency"],
        ["Vue + Babylon.js", "Dashboard multiskala", "Visual tapak, bangunan, dan indoor pada peramban"],
        ["JSON Schema + replay API", "Kontrak telemetry lokal", "Validasi dan distribusi energi–okupansi berprovenance"],
        ["ESP32 dan sensor legacy", "DHT11, ZMPT101B, SCT013", "Sumber historis trace kalibrasi; tidak diukur ulang"],
    ]
    add_caption(doc, "Tabel 3.3 Alat dan instrumen penelitian")
    add_simple_table(doc, ["Alat/instrumen", "Spesifikasi", "Kegunaan"], tool_rows, [3.4, 4.5, 6.1], font_size=9)

    subheading(doc, "3.3.3", "Variabel Penelitian", level=2)
    variable_rows = [
        ["Telemetry energi", "voltage_v, current_a, power_legacy, energy_cumulative_legacy_wh", "Daya legacy diverifikasi sebagai V×I; energi diintegrasikan per siklus"],
        ["Telemetry lingkungan/okupansi", "temperature_c, humidity_pct, people_count, occupancy_status", "Konteks bangunan sesuai data historis"],
        ["Konteks multiskala dan LoD", "application_lod, EPSG:4326, supported_views, scale_semantics", "Mengikat LoD-A tapak, LoD-B bangunan, dan LoD-C indoor; perpindahan manual"],
        ["Identitas sumber", "device_id, source_type, source_row_id, replay_block_id", "Keterlacakan pesan ke sumber dan blok replay"],
        ["Waktu dan status", "source_timestamp, replay_timestamp, route, validity_status", "Urutan, freshness, dan keputusan pemrosesan"],
        ["Metrik arsitektur", "P50/P95/P99, throughput, payload, routing, drop, deadline miss", "Karakteristik jalur edge–cloud"],
    ]
    add_caption(doc, "Tabel 3.4 Variabel penelitian")
    add_simple_table(doc, ["Kelompok", "Variabel", "Peran"], variable_rows, [3.0, 6.0, 5.0], font_size=8.5)

    subheading(doc, "3.3.4", "Rancangan Replay Data Historis", level=2)
    add_body(
        doc,
        "Workload 2.027.520 baris direkonstruksi menjadi 22 blok berukuran 92.160 baris. Setiap baris diberi replay_block_id dan source_row_id tanpa mengubah nilai payload. Urutan timestamp diperiksa, sedangkan API demonstrasi menerbitkan satu sampel sesuai interval polling 3,5 detik.",
    )
    add_body(
        doc,
        "Sebanyak 5.000 posisi dipilih secara merata dari seluruh blok untuk pengukuran latency per-pesan. Audit volume dan integritas tetap dilakukan terhadap keseluruhan dataset. Tidak ada pemisahan train, validation, atau test karena penelitian tidak melakukan pelatihan model.",
    )

    subheading(doc, "3.3.5", "Rancangan Evaluasi Arsitektur", level=2)
    add_body(
        doc,
        "Setiap pesan benchmark menjalani validasi skema dan nilai, pemetaan energi–okupansi, penambahan metadata provenance, serialisasi JSON, dan routing. Pesan dengan pembacaan listrik invalid, arus di bawah threshold legacy 0,1 A, atau daya legacy melebihi ambang P99 trace sebesar 42,6 W dirutekan ke cloud terkonfigurasi; pesan lain tetap pada edge.",
    )
    add_body(
        doc,
        "Waktu pemantauan lokal dan serialisasi diukur menggunakan monotonic high-resolution clock. Profil jaringan menggunakan median 45 ms, jitter 12 ms, dan probabilitas drop 0,005. Latensi end-to-end edge–cloud dibandingkan dengan baseline cloud-only pada draw jaringan yang sama. Seluruh angka jaringan diberi label emulated/configured.",
    )

    subheading(doc, "3.3.6", "Ancaman Validitas dan Mitigasi", level=2)
    threats = [
        ["Trace hanya satu gateway dan sekitar empat hari", "Membatasi klaim pada evaluasi monitoring satu sumber; tidak mengklaim banyak bangunan."],
        ["Tidak ada wattmeter atau faktor daya independen", "Melaporkan nilai sebagai daya legacy/proksi daya semu V×I, bukan daya aktif."],
        ["Energi merupakan integral proksi daya", "Membatasi hasil sebagai Wh legacy per siklus, bukan energi aktif terkalibrasi."],
        ["Okupansi berasal dari alur sensor legacy", "Menampilkan provenance dan tidak mengklaim validasi ground truth baru."],
        ["Workload legacy merupakan 22 replay", "Melaporkan volume sebagai workload dan mempertahankan pemetaan ke baris sumber."],
        ["Timestamp replay digeser", "Menyimpan timestamp sumber dan replay secara terpisah dalam provenance."],
        ["Benchmark bukan hardware edge produksi", "Melaporkan runtime mesin dan tidak menyebutnya sebagai Raspberry Pi."],
        ["Jaringan bukan public cloud nyata", "Memberi label emulasi dan memisahkan latensi terukur dari latensi terkonfigurasi."],
        ["Render browser belum masuk latency", "Menyatakan staleness sebagai proxy sampai pengukuran browser dilakukan."],
        ["Koordinat dan LoD geometrik belum divalidasi", "Menyatakan LoD-A/B/C sebagai hirarki aplikatif proyek; koordinat survei dan kepatuhan standar belum diuji."],
    ]
    add_caption(doc, "Tabel 3.5 Ancaman validitas dan mitigasi")
    add_simple_table(doc, ["Ancaman validitas", "Mitigasi"], threats, [6.5, 7.5], font_size=9)

    subheading(doc, "3.4", "Bukti Hasil Uji Pipeline")
    add_body(
        doc,
        "Sebagai bukti bahwa rancangan dapat dieksekusi, pipeline final telah dijalankan pada 2.027.520 baris workload dan mengambil 5.000 posisi yang mencakup seluruh 22 blok replay. Hasil ini merupakan benchmark arsitektur pada mesin eksperimen dan emulasi jaringan terkonfigurasi, bukan validasi lapangan baru.",
    )
    local = benchmark_results["actual_local_monitoring"]
    serialization = benchmark_results["actual_json_serialization"]
    edge_path = benchmark_results["actual_edge_path"]
    cloud_route = benchmark_results["configured_cloud_route_end_to_end"]
    end_to_end = benchmark_results["configured_end_to_end"]
    cloud_only = benchmark_results["configured_cloud_only_baseline"]
    architecture_comparison = benchmark_results["architecture_comparison"]
    result_rows = [
        ["Pemantauan lokal", f"{local['p50_ms']:.3f}", f"{local['p95_ms']:.3f}", f"{local['p99_ms']:.3f}", "Lokal terukur"],
        ["Serialisasi JSON", f"{serialization['p50_ms']:.3f}", f"{serialization['p95_ms']:.3f}", f"{serialization['p99_ms']:.3f}", "Lokal terukur"],
        ["Jalur edge", f"{edge_path['p50_ms']:.3f}", f"{edge_path['p95_ms']:.3f}", f"{edge_path['p99_ms']:.3f}", "Lokal terukur"],
        ["Jalur cloud", f"{cloud_route['p50_ms']:.3f}", f"{cloud_route['p95_ms']:.3f}", f"{cloud_route['p99_ms']:.3f}", "Skenario uji"],
        ["End-to-end campuran", f"{end_to_end['p50_ms']:.3f}", f"{end_to_end['p95_ms']:.3f}", f"{end_to_end['p99_ms']:.3f}", "Gabungan rute"],
        ["Baseline cloud-only", f"{cloud_only['p50_ms']:.3f}", f"{cloud_only['p95_ms']:.3f}", f"{cloud_only['p99_ms']:.3f}", "Kontrafaktual terkonfigurasi"],
    ]
    add_caption(doc, "Tabel 3.6 Ringkasan hasil latensi pipeline (ms)")
    add_simple_table(
        doc,
        ["Komponen", "P50", "P95", "P99", "Status"],
        result_rows,
        [3.2, 2.4, 2.4, 2.4, 3.6],
        font_size=8,
    )
    routing = benchmark_results["routing"]
    throughput = benchmark_results["throughput"]["sequential_messages_per_second"]
    add_body(
        doc,
        f"Benchmark menghasilkan throughput lokal {throughput:,.2f} pesan/detik, dengan {routing['edge_count']:,} pesan tetap pada edge dan {routing['cloud_count']:,} pesan dirutekan ke cloud terkonfigurasi. Dibanding baseline cloud-only pada kondisi emulasi yang sama, P95 turun {architecture_comparison['configured_p95_latency_reduction_percent']:.2f}% dan {architecture_comparison['network_payload_bytes_avoided']:,} byte transfer jaringan dihindari. Tidak terdapat deadline miss terhadap batas 3.500 ms. Visual berikut berasal langsung dari artefak hasil final pipeline.",
    )
    result_figures = [
        ("01_trace_profile.png", "Gambar 3.2 Profil trace historis asli"),
        ("02_replay_provenance.png", "Gambar 3.3 Cakupan dan provenance workload replay"),
        ("03_monitoring_checks.png", "Gambar 3.4 Hasil validasi pesan dan routing edge–cloud"),
        ("04_latency_characteristics.png", "Gambar 3.5 Ringkasan latensi dan routing edge–cloud"),
        ("05_multiscale_digital_twin.png", "Gambar 3.6 Pemetaan arsitektur dan Digital Twin multiskala"),
    ]
    for filename, caption in result_figures:
        add_figure(doc, FINAL_RESULTS / "figures" / filename, caption, 14.0)

    # References
    chapter(doc, "", "Referensi")
    # Remove empty BAB prefix produced by helper and rewrite current paragraph.
    last = doc.paragraphs[-1]
    last.clear()
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(last.add_run("REFERENSI"), bold=True)
    refs = [
        'M. N. Shahid, M. U. Shahid, and M. Irfan, “Advances in Building Energy Management: A Comprehensive Review,” Buildings, vol. 15, art. no. 4237, 2025, doi: 10.3390/buildings15234237.',
        'Y. Himeur, A. Alsalemi, F. Bensaali, and A. Amira, “The Emergence of Hybrid Edge-Cloud Computing for Energy Efficiency in Buildings,” Lecture Notes in Networks and Systems, vol. 295, pp. 70–83, 2022, doi: 10.1007/978-3-030-82196-8_6.',
        'M. Trigka and E. Dritsas, “Edge and Cloud Computing in Smart Cities,” Future Internet, vol. 17, art. no. 118, 2025, doi: 10.3390/fi17030118.',
        'I. Rojek, P. Prokopowicz, M. Piechowiak, P. Kotlarz, N. Náprstková, and D. Mikołajewski, “The Impact of Data Analytics Based on Internet of Things, Edge Computing, and Artificial Intelligence on Energy Efficiency in Smart Environment,” Applied Sciences, vol. 16, art. no. 225, 2026, doi: 10.3390/app16010225.',
        'A. Sghiri, M. Gallab, S. Merzouk, and S. Assoul, “Leveraging Digital Twins for Enhancing Building Energy Efficiency: A Literature Review of Applications, Technologies, and Challenges,” Buildings, vol. 15, art. no. 498, 2025, doi: 10.3390/buildings15030498.',
        'T. Sinthamrongruk, K. Dahal, and N. Harnpornchai, “A Lightweight Web3D Digital Twin Framework for Real-Time ESG Monitoring Using IoT Sensors,” Electronics, vol. 15, art. no. 1736, 2026, doi: 10.3390/electronics15081736.',
        'Z. Wang, H. Xiao, C. Guan, L. Zhou, and D. Fu, “Research on the Development of a Building Model Management System Integrating MQTT Sensing,” Sensors, vol. 25, art. no. 6069, 2025, doi: 10.3390/s25196069.',
        'S. Moveh, E. A. Merchán-Cruz, M. Abuhussain, Y. A. Dodo, S. Alhumaid, and A. H. Alhamami, “Deep Learning Framework Using Transformer Networks for Multi Building Energy Consumption Prediction in Smart Cities,” Energies, vol. 18, art. no. 1468, 2025, doi: 10.3390/en18061468.',
        'B. H. Suharto, S. H. Wijono, M. Hardiyanti, M. K. Fajarlestari, and D. L. Hakim, “Occupancy-Aware Spatio-Temporal Building Energy Forecasting with a Hybrid Long Short-Term Memory and Graph Neural Network Benchmark Using Public Datasets,” E3S Web of Conferences, vol. 687, art. no. 02006, 2026, doi: 10.1051/e3sconf/202668702006.',
        'K. Sun, “DMFF: Deep Multimodel Feature Fusion for Building Occupancy Detection,” Building and Environment, vol. 253, art. no. 111355, 2024, doi: 10.1016/j.buildenv.2024.111355.',
        'G. Zocchi, M. Hosseini, and G. Triantafyllidis, “Exploring the Synergy of Advanced Lighting Controls, Building Information Modelling and Internet of Things for Sustainable and Energy-Efficient Buildings: A Systematic Literature Review,” Sustainability, vol. 16, art. no. 10937, 2024, doi: 10.3390/su162410937.',
        'H. AlZaabi et al., “Intelligent Energy Consumption for Smart Homes Using Fused Machine-Learning Technique,” Computers, Materials & Continua, vol. 74, no. 1, pp. 2261–2278, 2023, doi: 10.32604/cmc.2023.031834.',
        'X. Luo et al., “Toward Intelligent AIoT: A Comprehensive Survey on Digital Twin and Multimodal Generative AI Integration,” Mathematics, vol. 13, art. no. 3382, 2025, doi: 10.3390/math13213382.',
        'Q. Chen, J. Chen, and W. Huang, “Visualizing Large-Scale Building Information Modeling Models within Indoor and Outdoor Environments Using a Semantics-Based Method,” ISPRS International Journal of Geo-Information, vol. 10, art. no. 756, 2021, doi: 10.3390/ijgi10110756.',
        'S. Herle, R. Becker, R. Wollenberg, and J. Blankenbach, “GIM and BIM: How to Obtain Interoperability Between Geospatial and Building Information Modelling?,” PFG—Journal of Photogrammetry, Remote Sensing and Geoinformation Science, vol. 88, pp. 33–42, 2020, doi: 10.1007/s41064-020-00090-4.',
        'A. Clausen et al., “A Digital Twin Framework for Improving Energy Efficiency and Occupant Comfort in Public and Commercial Buildings,” Energy Informatics, vol. 4, art. no. 40, 2021, doi: 10.1186/s42162-021-00153-9.',
        'I. Smirnov and F. Re Cecconi, “Occupancy-Aware Digital Twin for Sustainable Buildings,” Buildings, vol. 16, art. no. 1629, 2026, doi: 10.3390/buildings16081629.',
        'G. Walczyk and A. Ożadowicz, “Building Information Modeling and Digital Twins for Functional and Technical Design of Smart Buildings with Distributed IoT Networks—Review and New Challenges Discussion,” Future Internet, vol. 16, art. no. 225, 2024, doi: 10.3390/fi16070225.',
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        set_run_font(p.add_run(f"[{idx}] "), size=12)
        set_run_font(p.add_run(ref), size=12)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
